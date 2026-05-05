# Imports required libraries for file handling, UI, PDFs, DOCX, and LangChain
import io
import tempfile

import PyPDF2
import docx
import streamlit as st
from langchain.chains import RetrievalQA
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI

# Define embedding model and OpenRouter base URL
EMBEDDING_MODEL = "intfloat/e5-large-v2"
OPENROUTER_BASE = "https://openrouter.ai/api/v1"

# Prompt used for RAG (forces model to use only document context)
RAG_PROMPT = PromptTemplate(
    template=(
        "Use only the document content below to answer.\n\n"
        "{context}\n\n"
        "Question: {question}\n"
        "Answer:"
    ),
    input_variables=["context", "question"],
)

# Streamlit UI setup + session memory initialization
st.set_page_config(page_title="Document RAG Assistant", page_icon="📄", layout="centered")
st.title("📄 Document RAG Assistant")

st.session_state.setdefault("rag", None)
st.session_state.setdefault("history", [])

# -Read PDF file and extract text
def read_pdf(upload):
    reader = PyPDF2.PdfReader(io.BytesIO(upload.read()))
    return "".join(page.extract_text() or "" for page in reader.pages)

# Read DOCX file and extract text
def read_docx(upload):
    file = docx.Document(upload)
    return "\n".join(p.text for p in file.paragraphs)

# Detect file type and load text accordingly
def load_text(upload):
    if upload.name.endswith(".pdf"):
        return read_pdf(upload)
    if upload.name.endswith(".docx"):
        return read_docx(upload)
    return ""

# Split long text into overlapping chunks for embeddings
def chunk_text(text, size=500, overlap=100):
    chunks = []
    start = 0

    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start += size - overlap

    return [Document(page_content=c) for c in chunks]

# Load embedding model (cached for performance)
@st.cache_resource(show_spinner=True)
def load_embeddings():
    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

# -Build full RAG pipeline (vector DB + LLM + retriever)
def build_pipeline(docs, api_key, model, temp):
    embeddings = load_embeddings()

    vdb = Chroma.from_documents(
        documents=docs,
        embedding=embeddings,
        persist_directory=tempfile.mkdtemp(),
    )

    llm = ChatOpenAI(
        base_url=OPENROUTER_BASE,
        api_key=api_key,
        model=model,
        temperature=temp,
    )

    return RetrievalQA.from_chain_type(
        llm=llm,
        retriever=vdb.as_retriever(search_kwargs={"k": 5}),
        chain_type="stuff",
        chain_type_kwargs={"prompt": RAG_PROMPT},
    )


# Sidebar inputs (API key, model, temperature, file upload)
with st.sidebar:
    api_key = st.text_input("OpenRouter API Key", type="password")

    model = st.text_input("Model", value="openai/gpt-4o-mini")

    temperature = st.slider("Temperature", 0.0, 1.0, 0.1, 0.05)

    uploaded = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])


    # Build knowledge base button
    if st.button("Build Knowledge Base"):
        if not uploaded:
            st.error("Upload a file")
        elif not api_key:
            st.error("Enter API key")
        else:
            text = load_text(uploaded)
            docs = chunk_text(text)

            st.session_state.rag = build_pipeline(docs, api_key, model, temperature)
            st.session_state.history = []
            st.success("Ready!")

# Main chat UI section
st.divider()

if st.session_state.rag is None:
    st.info("Upload a document to start")

# Display previous chat history
else:
    for msg in st.session_state.history:
        st.chat_message(msg["role"]).markdown(msg["content"])

# Handle user question + RAG response
    if query := st.chat_input("Ask something about the document"):
        st.session_state.history.append({"role": "user", "content": query})
        st.chat_message("user").markdown(query)

        with st.chat_message("assistant"):
            result = st.session_state.rag.invoke({"query": query})
            answer = result.get("result", "")
            st.markdown(answer)

        st.session_state.history.append({"role": "assistant", "content": answer})

# Clear chat history
    if st.session_state.history and st.button("Clear chat"):
        st.session_state.history = []
        st.rerun()